from __future__ import annotations

from pathlib import Path

from docx import Document

from app.core.document_engine.contracts import DocumentMetadata, DocumentSource, ExtractedPage
from app.core.utils.common import extension_for, mime_type_for, normalize_text, sha256_bytes


def process_docx(path: Path, source: DocumentSource | None = None) -> tuple[DocumentMetadata, list[ExtractedPage]]:
    document = Document(str(path))
    paragraphs = [paragraph.text for paragraph in document.paragraphs if paragraph.text.strip()]
    tables = [" | ".join(cell.text.strip() for cell in row.cells) for table in document.tables for row in table.rows]
    text = normalize_text("\n".join([*paragraphs, *tables]))
    core_properties = document.core_properties
    metadata = DocumentMetadata(
        name=path.name,
        extension=extension_for(path.name),
        mime_type=mime_type_for(path.name),
        size_bytes=path.stat().st_size,
        file_hash=sha256_bytes(path.read_bytes()),
        page_count=1,
        author=core_properties.author or "",
        title=core_properties.title or "",
        created_at=core_properties.created,
        modified_at=core_properties.modified,
        extra={"paragraph_count": len(paragraphs), "table_count": len(document.tables)},
    )
    return metadata, [ExtractedPage(page_number=1, text=text)]
